from __future__ import annotations

from pathlib import Path

from ocr_book.config.schema import ExportConfig
from ocr_book.export.base import Exporter
from ocr_book.export.block_grouping import group_list_runs
from ocr_book.reconstruction.document_model import Block, BlockType, Document, TextRun
from ocr_book.utils.errors import EngineNotAvailableError


class DocxExporter(Exporter):
    extension = ".docx"

    def export(self, document: Document, output_path: Path, config: ExportConfig) -> Path:
        try:
            from docx import Document as DocxDocument
            from docx.shared import Inches, Pt
        except ImportError as exc:
            raise EngineNotAvailableError(
                "El exportador DOCX requiere python-docx. Instálalo con: pip install python-docx"
            ) from exc

        docx = DocxDocument()

        for item in group_list_runs(document.blocks):
            if isinstance(item, list):
                style = "List Number" if item[0].list_ordered else "List Bullet"
                for list_item in item:
                    paragraph = docx.add_paragraph(style=style)
                    self._add_runs(paragraph, list_item.runs)
            else:
                self._add_block(docx, item, Inches, Pt)

        output_path = output_path.with_suffix(self.extension)
        docx.save(str(output_path))
        return output_path

    @staticmethod
    def _add_runs(paragraph, runs: list[TextRun]) -> None:
        for run in runs:
            docx_run = paragraph.add_run(run.text)
            docx_run.bold = run.bold
            docx_run.italic = run.italic

    def _add_block(self, docx, block: Block, inches, pt) -> None:
        if block.type == BlockType.TITLE:
            paragraph = docx.add_paragraph(style="Title")
            self._add_runs(paragraph, block.runs)
        elif block.type == BlockType.HEADING:
            paragraph = docx.add_paragraph(style="Heading 1")
            self._add_runs(paragraph, block.runs)
        elif block.type == BlockType.QUOTE:
            paragraph = docx.add_paragraph(style="Intense Quote")
            self._add_runs(paragraph, block.runs)
        elif block.type == BlockType.FOOTNOTE:
            paragraph = docx.add_paragraph()
            if block.footnote_marker:
                marker_run = paragraph.add_run(f"{block.footnote_marker} ")
                marker_run.font.superscript = True
                marker_run.font.size = pt(9)
            self._add_runs(paragraph, block.runs)
            for run in paragraph.runs:
                run.font.size = pt(9)
        else:
            paragraph = docx.add_paragraph()
            self._add_runs(paragraph, block.runs)
            if block.first_line_indent:
                paragraph.paragraph_format.first_line_indent = inches(0.3)
